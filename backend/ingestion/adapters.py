"""Format adapters — extract text/tables only; parsing is separate."""

from __future__ import annotations

import csv
import io
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from typing import Any

from ingestion.normalize import sanitize_csv_formula
from ingestion.ocr.tesseract_engine import LocalOCRUnavailable, run_tesseract
from ingestion.signatures import IngestionReject


@dataclass
class AdapterResult:
    text: str
    rows: list[dict[str, Any]] = field(default_factory=list)
    page_texts: list[str] = field(default_factory=list)
    source_kind: str = "UNKNOWN"
    page_count: int | None = None
    sheet_names: list[str] = field(default_factory=list)
    ocr_confidence: float | None = None
    complex_layout: bool = False
    warnings: list[str] = field(default_factory=list)
    adapter: str = ""
    needs_local_ocr: bool = False


def extract_xlsx(content: bytes) -> AdapterResult:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise IngestionReject(
            "OCR_NOT_CONFIGURED",
            "openpyxl is not installed for XLSX ingestion.",
        ) from exc

    bio = io.BytesIO(content)
    try:
        wb = load_workbook(bio, read_only=True, data_only=True)
    except Exception as exc:
        raise IngestionReject("UNSUPPORTED_DOCUMENT", f"Malformed XLSX: {exc}") from exc

    sheets = list(wb.sheetnames)
    rows: list[dict[str, Any]] = []
    text_parts: list[str] = []
    warnings: list[str] = []

    for sheet_name in sheets:
        ws = wb[sheet_name]
        headers: list[str] | None = None
        for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [("" if cell is None else str(cell).strip()) for cell in row]
            if not any(values):
                continue
            # formula injection markers in cached values
            for cell in values:
                if cell.startswith(("=", "+", "-", "@")):
                    warnings.append(
                        f"Sheet {sheet_name} row {row_index} contains formula-like text; preserved literally."
                    )
            if headers is None:
                headers = [v or f"col_{i}" for i, v in enumerate(values)]
                text_parts.append(f"# Sheet: {sheet_name}")
                text_parts.append(" | ".join(headers))
                continue
            record = {headers[i]: values[i] if i < len(values) else "" for i in range(len(headers))}
            rows.append(record)
            text_parts.append(" | ".join(values))
            if len(rows) > 50000:
                warnings.append("Row cap reached (50000); remaining rows skipped.")
                break

    wb.close()
    if not rows and not text_parts:
        raise IngestionReject("DOCUMENT_UNREADABLE", "XLSX contained no readable rows.")
    return AdapterResult(
        text="\n".join(text_parts),
        rows=rows,
        source_kind="XLSX",
        sheet_names=sheets,
        warnings=warnings,
        adapter="openpyxl-readonly",
        complex_layout=len(sheets) > 1,
    )


def _detect_csv_dialect(sample: str) -> csv.Dialect:
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        class Fallback(csv.Dialect):
            delimiter = ","
            quotechar = '"'
            escapechar = None
            doublequote = True
            skipinitialspace = True
            lineterminator = "\n"
            quoting = csv.QUOTE_MINIMAL

        return Fallback()


def extract_csv(content: bytes) -> AdapterResult:
    encoding = "utf-8"
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            import chardet

            guess = chardet.detect(content) or {}
            encoding = guess.get("encoding") or "latin-1"
            text = content.decode(encoding, errors="replace")
        except Exception:
            text = content.decode("latin-1", errors="replace")

    sample = text[:4096]
    dialect = _detect_csv_dialect(sample)
    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    rows: list[dict[str, Any]] = []
    warnings: list[str] = []
    for row in reader:
        cleaned = {k: sanitize_csv_formula(v) for k, v in (row or {}).items() if k}
        for value in cleaned.values():
            if str(value).startswith("'=") or str(value).startswith("="):
                warnings.append("Formula-like CSV cell neutralized for safe reporting.")
        rows.append(cleaned)
        if len(rows) > 100000:
            warnings.append("CSV row cap reached.")
            break
    if not rows:
        raise IngestionReject("DOCUMENT_UNREADABLE", "CSV contained no data rows.")
    # rebuild text from rows for parser
    headers = list(rows[0].keys())
    lines = [" | ".join(headers)]
    for row in rows[:2000]:
        lines.append(" | ".join(str(row.get(h, "")) for h in headers))
    return AdapterResult(
        text="\n".join(lines),
        rows=rows,
        source_kind="CSV",
        warnings=warnings,
        adapter=f"csv/{encoding}",
    )


def _pdf_is_encrypted(doc: Any) -> bool:
    try:
        return bool(doc.is_encrypted)
    except Exception:
        return False


def extract_pdf(content: bytes, *, language_hint: str = "UNKNOWN") -> AdapterResult:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise IngestionReject(
            "OCR_NOT_CONFIGURED",
            "PyMuPDF is not installed for PDF ingestion.",
        ) from exc

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise IngestionReject("UNSUPPORTED_DOCUMENT", f"Malformed PDF: {exc}") from exc

    if _pdf_is_encrypted(doc):
        # try empty password; otherwise reject
        try:
            auth = doc.authenticate("")
        except Exception:
            auth = 0
        if not auth:
            doc.close()
            raise IngestionReject(
                "UNSUPPORTED_DOCUMENT",
                "Password-protected PDF is not supported without an unlock workflow.",
            )

    max_pages = max(1, int(os.getenv("INGESTION_MAX_PDF_PAGES", "30")))
    if doc.page_count > max_pages:
        doc.close()
        raise IngestionReject(
            "PDF_PAGE_LIMIT_EXCEEDED",
            f"PDF exceeds configured page limit ({max_pages}).",
        )

    page_texts: list[str] = []
    digital_flags: list[bool] = []
    warnings: list[str] = []
    ocr_scores: list[float] = []
    needs_ocr_pages = 0

    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        text = (page.get_text("text") or "").strip()
        usable = len(re.sub(r"\s+", "", text)) >= 40
        digital_flags.append(usable)
        if usable:
            page_texts.append(text)
            continue
        needs_ocr_pages += 1
        # render controlled resolution for local OCR
        zoom = 2.0
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image_bytes = pix.tobytes("png")
        try:
            ocr = run_tesseract(image_bytes, language_hint=language_hint, layout_hint="TABLE")
            page_texts.append(ocr.text)
            ocr_scores.append(ocr.confidence)
            warnings.extend(ocr.warnings)
        except LocalOCRUnavailable as exc:
            page_texts.append("")
            warnings.append(f"Page {page_index + 1}: {exc.message}")

    doc.close()
    digital_count = sum(1 for flag in digital_flags if flag)
    scanned_count = len(digital_flags) - digital_count
    if digital_count and scanned_count:
        kind = "PDF_MIXED"
    elif scanned_count:
        kind = "PDF_SCANNED"
    else:
        kind = "PDF_DIGITAL"

    combined = "\n\n".join(t for t in page_texts if t)
    if not combined.strip():
        raise IngestionReject("DOCUMENT_UNREADABLE", "PDF produced no readable text.")

    return AdapterResult(
        text=combined,
        page_texts=page_texts,
        source_kind=kind,
        page_count=len(page_texts),
        ocr_confidence=(sum(ocr_scores) / len(ocr_scores)) if ocr_scores else None,
        complex_layout=True,
        warnings=warnings,
        adapter="pymupdf",
        needs_local_ocr=needs_ocr_pages > 0,
        rows=_tables_from_text_blocks(combined),
    )


def _tables_from_text_blocks(text: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line in text.splitlines():
        if "|" in line:
            parts = [p.strip() for p in line.split("|") if p.strip()]
        elif "\t" in line:
            parts = [p.strip() for p in line.split("\t") if p.strip()]
        else:
            continue
        if len(parts) >= 2 and any(ch.isdigit() for ch in "".join(parts)):
            record: dict[str, Any] = {"month": parts[0], "amount": parts[1]}
            if len(parts) > 2:
                record["col2"] = parts[2]
            rows.append(record)
    return rows


def extract_docx(content: bytes) -> AdapterResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestionReject(
            "OCR_NOT_CONFIGURED",
            "python-docx is not installed for DOCX ingestion.",
        ) from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise IngestionReject("UNSUPPORTED_DOCUMENT", f"Malformed DOCX: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    rows: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        headers: list[str] | None = None
        for row in table.rows:
            values = [(cell.text or "").strip() for cell in row.cells]
            if not any(values):
                continue
            if headers is None:
                headers = [v or f"col_{i}" for i, v in enumerate(values)]
                parts.append(" | ".join(headers))
                continue
            record = {headers[i]: values[i] if i < len(values) else "" for i in range(len(headers))}
            rows.append(record)
            parts.append(" | ".join(values))
        parts.append(f"[table {table_index + 1} end]")

    text = "\n".join(parts)
    if not text.strip():
        raise IngestionReject("DOCUMENT_UNREADABLE", "DOCX contained no paragraphs or tables.")
    return AdapterResult(
        text=text,
        rows=rows,
        source_kind="DOCX",
        adapter="python-docx",
        complex_layout=bool(rows),
    )


def _libreoffice_bin() -> str | None:
    configured = os.getenv("INGESTION_LIBREOFFICE_BIN", "").strip()
    if configured:
        return configured
    return shutil.which("soffice") or shutil.which("libreoffice")


def extract_doc(content: bytes) -> AdapterResult:
    binary = _libreoffice_bin()
    if not binary:
        raise IngestionReject(
            "OCR_NOT_CONFIGURED",
            "LibreOffice is not available for legacy DOC conversion.",
        )
    timeout = max(5, int(os.getenv("INGESTION_LIBREOFFICE_TIMEOUT_SECONDS", "60")))
    max_mb = max(1, float(os.getenv("INGESTION_MAX_FILE_MB", "15")))
    if len(content) > max_mb * 1024 * 1024:
        raise IngestionReject("FILE_TOO_LARGE", "DOC exceeds configured size limit.")

    temp_dir = tempfile.mkdtemp(prefix="vardhan-doc-")
    try:
        source_path = os.path.join(temp_dir, "input.doc")
        with open(source_path, "wb") as handle:
            handle.write(content)
        try:
            subprocess.run(
                [
                    binary,
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    source_path,
                ],
                check=True,
                timeout=timeout,
                capture_output=True,
            )
        except subprocess.TimeoutExpired as exc:
            raise IngestionReject("OCR_TIMEOUT", "LibreOffice DOC conversion timed out.") from exc
        except subprocess.CalledProcessError as exc:
            # fallback to PDF conversion
            try:
                subprocess.run(
                    [
                        binary,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        temp_dir,
                        source_path,
                    ],
                    check=True,
                    timeout=timeout,
                    capture_output=True,
                )
            except Exception as pdf_exc:
                raise IngestionReject(
                    "OCR_FAILED",
                    f"LibreOffice failed to convert DOC: {exc}",
                ) from pdf_exc

        docx_path = os.path.join(temp_dir, "input.docx")
        pdf_path = os.path.join(temp_dir, "input.pdf")
        if os.path.isfile(docx_path):
            with open(docx_path, "rb") as handle:
                converted = handle.read()
            result = extract_docx(converted)
            result.source_kind = "DOC"
            result.adapter = "libreoffice->docx"
            return result
        if os.path.isfile(pdf_path):
            with open(pdf_path, "rb") as handle:
                converted = handle.read()
            result = extract_pdf(converted)
            result.source_kind = "DOC"
            result.adapter = "libreoffice->pdf"
            return result
        raise IngestionReject("OCR_FAILED", "LibreOffice conversion produced no output file.")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def extract_image(content: bytes, *, language_hint: str = "UNKNOWN") -> AdapterResult:
    try:
        ocr = run_tesseract(content, language_hint=language_hint, layout_hint="REGISTER")
    except LocalOCRUnavailable as exc:
        raise IngestionReject("OCR_NOT_CONFIGURED", exc.message) from exc
    if not ocr.text.strip():
        raise IngestionReject("DOCUMENT_UNREADABLE", "Image OCR produced no text.")
    return AdapterResult(
        text=ocr.text,
        page_texts=[ocr.text],
        source_kind="IMAGE",
        page_count=1,
        ocr_confidence=ocr.confidence,
        warnings=ocr.warnings,
        adapter="tesseract",
        needs_local_ocr=True,
        rows=_tables_from_text_blocks(ocr.text),
    )
