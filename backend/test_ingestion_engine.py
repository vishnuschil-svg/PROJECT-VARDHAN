"""Targeted tests for the Universal File Ingestion Engine."""

from __future__ import annotations

import hashlib
import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import AsyncMock, patch

from ingestion.escalation import merge_gemini_into_draft, should_escalate
from ingestion.file_router import route_and_extract
from ingestion.normalize import parse_indian_date, parse_money, sanitize_csv_formula
from ingestion.parser.chit_parser import parse_extracted_content
from ingestion.queue.store import IngestionQueueStore
from ingestion.schemas import ChitPlanDraft, empty_draft
from ingestion.service import IngestionService
from ingestion.signatures import IngestionReject, detect_file


FIXTURES = Path(__file__).resolve().parent / "fixtures" / "ingestion"


def _write_xlsx(path: Path, sheets: dict[str, list[list[object]]]) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    first = True
    for name, rows in sheets.items():
        ws = wb.active if first else wb.create_sheet(name)
        if first:
            ws.title = name
            first = False
        for row in rows:
            ws.append(row)
    wb.save(path)


def _write_docx(path: Path, paragraphs: list[str], table: list[list[str]] | None = None) -> None:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table:
        tbl = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                tbl.cell(r, c).text = value
    doc.save(path)


def _write_pdf_text(path: Path, text: str) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def ensure_fixtures() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    valid = FIXTURES / "valid_chit.xlsx"
    if not valid.exists():
        _write_xlsx(
            valid,
            {
                "Plan": [
                    ["Field", "Value"],
                    ["Chit Name", "Sri Rama Chit"],
                    ["Chit Value", "100000"],
                    ["Member Count", "20"],
                    ["Duration", "20"],
                    ["Monthly Installment", "5000"],
                    ["Installment Pattern", "Fixed Monthly"],
                ]
            },
        )
    multi = FIXTURES / "multi_sheet.xlsx"
    if not multi.exists():
        _write_xlsx(
            multi,
            {
                "Plan": [
                    ["Chit Name", "Multi Sheet Chit"],
                    ["Chit Value", "200000"],
                    ["Members", "25"],
                    ["Months", "25"],
                ],
                "Schedule": [
                    ["Month", "Installment"],
                    [1, 8000],
                    [2, 8000],
                ],
            },
        )
    partial = FIXTURES / "partial_invalid.xlsx"
    if not partial.exists():
        _write_xlsx(
            partial,
            {
                "Plan": [
                    ["Chit Name", "Partial Chit"],
                    ["Notes", "value missing on purpose"],
                ]
            },
        )
    csv_path = FIXTURES / "plan.csv"
    if not csv_path.exists():
        csv_path.write_text(
            "Chit Name,Chit Value,Members,Months,Installment\n"
            "CSV Chit,\"1,00,000\",20,20,5000\n",
            encoding="utf-8",
        )
    pdf_path = FIXTURES / "digital_plan.pdf"
    if not pdf_path.exists():
        _write_pdf_text(
            pdf_path,
            "Chit Name: Digital PDF Chit\nChit Value: 150000\nMembers: 15\nDuration: 15\n"
            "Monthly Installment: 10000\nFixed Monthly installment\n",
        )
    docx_path = FIXTURES / "plan.docx"
    if not docx_path.exists():
        _write_docx(
            docx_path,
            [
                "Chit Name: Docx Plan Chit",
                "Chit Value: 120000",
                "Member Count: 12",
                "Duration: 12",
                "Monthly Installment: 10000",
                "Terms and Conditions: Pay before the 5th.",
            ],
            table=[
                ["Month", "Installment"],
                ["1", "10000"],
                ["2", "10000"],
            ],
        )
    eng = FIXTURES / "english_plan.txt"
    if not eng.exists():
        eng.write_text(
            "Chit Name: English Image Chit\nChit Value: ₹2,00,000\nMembers: 20\n"
            "Duration: 20 months\nMonthly Installment: 10000\nAuction rules apply.\n",
            encoding="utf-8",
        )
    conflict = FIXTURES / "conflict.txt"
    if not conflict.exists():
        conflict.write_text(
            "Chit Name: Conflict Chit\nChit Value: 999999\nMembers: 10\n"
            "Duration: 20\nMonthly Installment: 1000\n",
            encoding="utf-8",
        )
    missing = FIXTURES / "missing.txt"
    if not missing.exists():
        missing.write_text("Chit Name: Missing Only\n", encoding="utf-8")


class NormalizeTests(unittest.TestCase):
    def test_money_and_dates(self):
        self.assertEqual(parse_money("₹1,00,000"), 100000.0)
        self.assertIsNone(parse_money("0"))
        self.assertIsNone(parse_money(""))
        self.assertEqual(parse_indian_date("15/01/2026"), "2026-01-15")
        self.assertEqual(sanitize_csv_formula("=cmd"), "'=cmd")


class SignatureTests(unittest.TestCase):
    def test_reject_executable(self):
        with self.assertRaises(IngestionReject):
            detect_file(b"MZ\x90\x00fake", "application/octet-stream", "virus.exe")

    def test_csv_detection(self):
        detected = detect_file(b"a,b\n1,2\n", "text/csv", "plan.csv")
        self.assertEqual(detected.source_kind_hint, "CSV")


class AdapterAndParserTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        ensure_fixtures()

    def test_valid_xlsx(self):
        content = (FIXTURES / "valid_chit.xlsx").read_bytes()
        routed = route_and_extract(content, filename="valid_chit.xlsx", declared_mime="")
        self.assertEqual(routed.result.source_kind, "XLSX")
        draft = parse_extracted_content(empty_draft(file_name="valid_chit.xlsx", mime_type=routed.detected.mime_type), text=routed.result.text, rows=routed.result.rows)
        self.assertEqual(draft.plan.chitName, "Sri Rama Chit")
        self.assertEqual(draft.plan.chitValue, 100000.0)
        self.assertEqual(draft.plan.memberCount, 20)

    def test_multi_sheet_xlsx(self):
        content = (FIXTURES / "multi_sheet.xlsx").read_bytes()
        routed = route_and_extract(content, filename="multi_sheet.xlsx", declared_mime="")
        self.assertGreaterEqual(len(routed.result.sheet_names), 2)
        self.assertTrue(routed.result.complex_layout)

    def test_partial_xlsx_missing_fields(self):
        content = (FIXTURES / "partial_invalid.xlsx").read_bytes()
        routed = route_and_extract(content, filename="partial_invalid.xlsx", declared_mime="")
        draft = parse_extracted_content(empty_draft(file_name="partial.xlsx", mime_type=routed.detected.mime_type), text=routed.result.text, rows=routed.result.rows)
        self.assertIn("plan.chitValue", draft.review.missingMandatoryFields)

    def test_csv(self):
        content = (FIXTURES / "plan.csv").read_bytes()
        routed = route_and_extract(content, filename="plan.csv", declared_mime="text/csv")
        draft = parse_extracted_content(empty_draft(file_name="plan.csv", mime_type="text/csv"), text=routed.result.text, rows=routed.result.rows)
        self.assertEqual(draft.plan.chitName, "CSV Chit")
        self.assertEqual(draft.plan.chitValue, 100000.0)

    def test_digital_pdf(self):
        content = (FIXTURES / "digital_plan.pdf").read_bytes()
        routed = route_and_extract(content, filename="digital_plan.pdf", declared_mime="application/pdf")
        self.assertEqual(routed.result.source_kind, "PDF_DIGITAL")
        draft = parse_extracted_content(empty_draft(file_name="digital_plan.pdf", mime_type="application/pdf"), text=routed.result.text)
        self.assertEqual(draft.plan.chitName, "Digital PDF Chit")

    def test_docx_paragraphs_and_tables(self):
        content = (FIXTURES / "plan.docx").read_bytes()
        routed = route_and_extract(content, filename="plan.docx", declared_mime="")
        self.assertEqual(routed.result.source_kind, "DOCX")
        self.assertTrue(routed.result.rows)
        draft = parse_extracted_content(empty_draft(file_name="plan.docx", mime_type=routed.detected.mime_type), text=routed.result.text, rows=routed.result.rows)
        self.assertEqual(draft.plan.chitName, "Docx Plan Chit")
        self.assertTrue(draft.terms)

    def test_missing_and_conflict_text(self):
        missing = (FIXTURES / "missing.txt").read_text(encoding="utf-8")
        draft = parse_extracted_content(empty_draft(file_name="missing.txt", mime_type="text/plain"), text=missing)
        self.assertIn("plan.chitValue", draft.review.missingMandatoryFields)

        conflict = (FIXTURES / "conflict.txt").read_text(encoding="utf-8")
        draft2 = parse_extracted_content(empty_draft(file_name="conflict.txt", mime_type="text/plain"), text=conflict)
        self.assertIn("plan.chitValue", draft2.review.conflictingFields)

    def test_never_invent_zero(self):
        draft = parse_extracted_content(
            empty_draft(file_name="z.txt", mime_type="text/plain"),
            text="Chit Name: Zero Trap\nChit Value: 0\nMembers: 0\n",
        )
        self.assertIsNone(draft.plan.chitValue)
        self.assertIsNone(draft.plan.memberCount)


class QueueAndServiceTests(unittest.IsolatedAsyncioTestCase):
    @classmethod
    def setUpClass(cls):
        ensure_fixtures()

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.store = IngestionQueueStore(Path(self.tmp.name) / "queue.db")
        self.service = IngestionService(self.store)

    def tearDown(self):
        self.tmp.cleanup()

    def test_duplicate_upload(self):
        content = (FIXTURES / "plan.csv").read_bytes()
        first = self.service.enqueue(
            content=content,
            filename="plan.csv",
            declared_mime="text/csv",
            tenant_id="tenant-a",
            workspace_id="ws-a",
            user_id="user-a",
        )
        second = self.service.enqueue(
            content=content,
            filename="plan.csv",
            declared_mime="text/csv",
            tenant_id="tenant-a",
            workspace_id="ws-a",
            user_id="user-a",
        )
        self.assertTrue(second.get("deduplicated"))
        self.assertEqual(first["id"], second["id"])
        self.assertEqual(hashlib.sha256(content).hexdigest(), first["sha256"])

    def test_gemini_unavailable_continues_with_local(self):
        content = (FIXTURES / "digital_plan.pdf").read_bytes()
        from vision_providers import VisionProviderError

        with patch(
            "ingestion.service.run_gemini_escalation",
            new=AsyncMock(side_effect=VisionProviderError("OCR_NOT_CONFIGURED", "not configured")),
        ), patch(
            "ingestion.service.should_escalate",
            return_value=(True, "forced"),
        ):
            job = self.service.enqueue(
                content=content,
                filename="digital_plan.pdf",
                declared_mime="application/pdf",
                tenant_id="tenant-b",
                workspace_id="ws-b",
                user_id="user-b",
            )
        self.assertEqual(job["status"], "NEEDS_REVIEW")
        self.assertIsNotNone(job.get("draft"))
        self.assertTrue(any("OCR_NOT_CONFIGURED" in w for w in job["draft"]["review"]["warnings"]))

    def test_gemini_rate_limited_preserves_local(self):
        content = (FIXTURES / "digital_plan.pdf").read_bytes()
        from vision_providers import VisionProviderError

        with patch(
            "ingestion.service.run_gemini_escalation",
            new=AsyncMock(side_effect=VisionProviderError("OCR_RATE_LIMIT", "quota", retryable=False)),
        ), patch(
            "ingestion.service.should_escalate",
            return_value=(True, "forced"),
        ):
            job = self.service.enqueue(
                content=content,
                filename="digital_plan.pdf",
                declared_mime="application/pdf",
                tenant_id="tenant-c",
                workspace_id="ws-c",
                user_id="user-c",
            )
        self.assertEqual(job["status"], "RATE_LIMITED")
        self.assertEqual(job.get("error_code"), "OCR_RATE_LIMIT")
        self.assertEqual(job["draft"]["plan"]["chitName"], "Digital PDF Chit")

    def test_unreadable_document(self):
        with self.assertRaises(IngestionReject):
            route_and_extract(b"\x00\x01\x02\x03not-a-file", filename="x.bin", declared_mime="application/octet-stream")

    def test_confirm_requires_mandatory_fields(self):
        content = (FIXTURES / "missing.txt").read_bytes()
        # treat as csv-like text via forced csv adapter path using labeled text file renamed
        job = self.store.create_job(
            tenant_id="t",
            workspace_id="w",
            user_id="u",
            file_name="missing.txt",
            mime_type="text/plain",
            sha256=hashlib.sha256(content).hexdigest(),
            byte_size=len(content),
        )
        draft = parse_extracted_content(empty_draft(file_name="missing.txt", mime_type="text/plain"), text=content.decode("utf-8"))
        self.store.update_job(job["id"], status="NEEDS_REVIEW", draft_json=draft.model_dump_json())
        with self.assertRaises(IngestionReject):
            self.service.confirm(job["id"], editor="organizer")

    def test_merge_gemini_does_not_overwrite_with_empty(self):
        draft = parse_extracted_content(
            empty_draft(file_name="x.txt", mime_type="text/plain"),
            text="Chit Name: Keep Me\nChit Value: 50000\nMembers: 10\nDuration: 10\nMonthly Installment: 5000\n",
        )

        class Extraction:
            chitName = None
            chitValue = None
            memberCount = None
            durationMonths = None
            monthlyInstallment = None
            organizerName = "Gemini Org"
            installmentPattern = "UNKNOWN"
            installmentSchedule = []
            startDate = None
            contactNumber = None
            foremanCommissionPercent = None
            minimumDiscountPercent = None
            maximumDiscountPercent = None
            auctionPattern = None
            specialRules = None
            notes = None

        class Result:
            extraction = Extraction()
            rawText = ""

        merged = merge_gemini_into_draft(draft, Result())
        self.assertEqual(merged.plan.chitName, "Keep Me")
        self.assertEqual(merged.plan.organizerName, "Gemini Org")


class EscalationGateTests(unittest.TestCase):
    def test_force_and_missing(self):
        draft = empty_draft(file_name="a", mime_type="image/jpeg")
        draft.review.missingMandatoryFields = ["plan.chitValue"]
        ok, reason = should_escalate(draft, complex_layout=False, ocr_confidence=0.9)
        self.assertTrue(ok)
        self.assertEqual(reason, "missing_required_fields")
        ok2, reason2 = should_escalate(draft, complex_layout=False, ocr_confidence=0.9, force=True)
        self.assertEqual(reason2, "user_ai_reprocess")


class OptionalBinaryTests(unittest.TestCase):
    def test_legacy_doc_without_libreoffice(self):
        from ingestion import adapters

        if adapters._libreoffice_bin():
            self.skipTest("LibreOffice present; skip negative path")
        with self.assertRaises(IngestionReject) as ctx:
            adapters.extract_doc(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64)
        self.assertEqual(ctx.exception.code, "OCR_NOT_CONFIGURED")

    def test_image_without_tesseract_or_with(self):
        # minimal valid 1x1 PNG
        png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00"
            b"\x00\x01\x01\x00\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        from ingestion.ocr.tesseract_engine import is_tesseract_available
        from ingestion import adapters

        if not is_tesseract_available():
            with self.assertRaises(IngestionReject) as ctx:
                adapters.extract_image(png)
            self.assertEqual(ctx.exception.code, "OCR_NOT_CONFIGURED")
        else:
            # May be unreadable due to tiny image; either text or DOCUMENT_UNREADABLE
            try:
                result = adapters.extract_image(png)
                self.assertEqual(result.source_kind, "IMAGE")
            except IngestionReject as exc:
                self.assertIn(exc.code, {"DOCUMENT_UNREADABLE", "OCR_NOT_CONFIGURED"})


if __name__ == "__main__":
    ensure_fixtures()
    unittest.main()
